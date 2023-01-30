# routes/ocr.py
from fastapi import APIRouter, UploadFile, Depends
from fastapi.responses import JSONResponse,  StreamingResponse
from typing import List
from app import oauth2
import pymongo
import os
import cv2
import io
import numpy as np
import pytesseract
from datetime import datetime
from app.database import OCR
from app.config import settings
import fitz
import pandas as pd
from docx import Document
import boto3


ocr = APIRouter()

s3 = boto3.client(
    "s3",
    aws_access_key_id=settings.AWS_ACCESS_KEY,
    aws_secret_access_key=settings.AWS_SECRET_KEY,
    region_name=settings.AWS_REGION
)


def read_pdf(pdf_file):
    pdf = fitz.open(stream=pdf_file)
    text = ""
    for page in pdf:
        text += page.get_text("text")
    return text


def read_excel(excel_file):
    df = pd.read_excel(excel_file)
    text = "".join(df.to_string())
    return text


def read_csv(csv_file):
    df = pd.read_csv(csv_file)
    text = "".join(df.to_string())
    return text


def read_word(word_file):
    doc = Document(word_file)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += '\n'+cell.text
    return text


def read_img(img):
    text = pytesseract.image_to_string(img)
    return (text)


@ocr.post("/extract_text")
async def extract_text(files: List[UploadFile], user_id: int = Depends(oauth2.require_user)):
    extracted_texts = []
    for file in files:
        text = ""
        if file.content_type == "application/pdf":
            pdf_file = io.BytesIO(await file.read())
            text = read_pdf(pdf_file)
        elif file.content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            excel_file = io.BytesIO(await file.read())
            text = read_excel(excel_file)
        elif file.content_type == "text/csv":
            csv_file = io.BytesIO(await file.read())
            text = read_csv(csv_file)
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            word_file = io.BytesIO(await file.read())
            text = read_word(word_file)
        else:
            img = await file.read()
            image_stream = io.BytesIO(img)
            image_stream.seek(0)
            file_bytes = np.asarray(
                bytearray(image_stream.read()), dtype=np.uint8)
            frame = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
            text = read_img(frame)

        extracted_texts.append({
            "user_id": user_id,
            "text": text,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })

        filename, file_extension = os.path.splitext(file.filename)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H:%M:%S")
        s3_key = f"{user_id}/{filename}_{timestamp}{file_extension}"
        s3.put_object(Bucket=settings.AWS_BUCKET_NAME,
                      Key=s3_key, Body=text)
    OCR.update_many(
        {}, {"$push": {"data": {"$each": extracted_texts}}}, upsert=True)
    return JSONResponse(content={"extracted_texts": extracted_texts})


@ ocr.get("/list_files")
async def list_files(user_id: int = Depends(oauth2.require_user)):
    try:
        # Use the boto3 client to list the objects in the S3 bucket
        objects = s3.list_objects(
            Bucket=settings.AWS_BUCKET_NAME, Prefix=f"{user_id}/")
        if not objects.get("Contents"):
            return JSONResponse(content={"status": "error", "message": "Data not found"}, status_code=404)
        # Extract the file names and sizes from the response
        file_list = []
        for obj in objects["Contents"]:
            filename = obj["Key"].replace(f"{user_id}/", "")
            filesize = obj["Size"]
            file_list.append({"name": filename, "size": filesize})
        return JSONResponse(content={"status": "success", "file_list": file_list})
    except Exception as e:
        return JSONResponse(content={"status": "error", "message": str(e)}, status_code=500)


@ ocr.get("/download_file/{file_name}")
async def download_file(file_name: str, user_id: int = Depends(oauth2.require_user)):
    try:
        # Use the boto3 client to check if the object exists in the S3 bucket
        s3.head_object(Bucket=settings.AWS_BUCKET_NAME,
                       Key=f"{user_id}/{file_name}")
        # Use the boto3 client to download the object from the S3 bucket
        file_content = s3.get_object(
            Bucket=settings.AWS_BUCKET_NAME, Key=f"{user_id}/{file_name}")["Body"].read()
        # Return the file content as an attachment
        file = io.BytesIO(file_content)
        return StreamingResponse(file, media_type="text/plain", headers={"Content-Disposition": f"attachment;filename={file_name}"})
    except Exception as e:
        if "NoSuchKey" in str(e):
            return JSONResponse(content={"status": "error", "message": "File not found"}, status_code=404)
        else:
            return JSONResponse(content={"status": "error", "message": str(e)}, status_code=500)
